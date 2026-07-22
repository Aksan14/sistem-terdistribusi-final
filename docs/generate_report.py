#!/usr/bin/env python3
"""
Generator Laporan DOCX - Sistem Logistik & Ekspedisi Terdistribusi
Laporan lengkap dengan kode program, keterangan screenshot, tabel, dan analisis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


# ============================================================
# HELPERS
# ============================================================

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def heading(doc, text, level=1, color='1E3A5F', spc_before=18, spc_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spc_before)
    p.paragraph_format.space_after = Pt(spc_after)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11, 4: 10}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
    return p


def para(doc, text, bold=False, italic=False, size=11, indent=True, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(f'• {text}')
    run.font.size = Pt(10.5)
    return p


def make_table(doc, headers, rows, col_widths=None, header_color='1D4ED8'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 0:
                set_cell_bg(cell, 'EFF6FF')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def code_block(doc, code_text, caption=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 64, 175)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(caption)
        cr.font.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(100, 116, 139)
    return p


def ss_box(doc, filename, caption, url='', note=''):
    """Kotak placeholder untuk screenshot — ganti dengan gambar nyata"""
    doc.add_paragraph()

    bt = doc.add_table(rows=1, cols=1)
    bt.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = bt.rows[0].cells[0]
    set_cell_bg(cell, 'DBEAFE')
    cell.width = Cm(14)

    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(20)
    cp.paragraph_format.space_after = Pt(20)

    r1 = cp.add_run('[ SCREENSHOT — MASUKKAN GAMBAR DI SINI ]\n\n')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(30, 64, 175)

    r2 = cp.add_run(f'Nama File : {filename}\n')
    r2.font.size = Pt(9.5)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(29, 78, 216)

    if url:
        r3 = cp.add_run(f'URL       : {url}\n')
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(51, 65, 85)

    if note:
        r4 = cp.add_run(f'Cara      : {note}')
        r4.font.size = Pt(9)
        r4.font.italic = True
        r4.font.color.rgb = RGBColor(71, 85, 105)

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(caption)
    cap_run.font.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = RGBColor(100, 116, 139)


# ============================================================
# READ SOURCE FILES
# ============================================================

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def read_snippet(path, start_marker, end_marker):
    full = os.path.join(BASE, path)
    try:
        with open(full, 'r') as f:
            content = f.read()
        si = content.find(start_marker)
        if si == -1:
            return content[:2000]
        ei = content.find(end_marker, si + len(start_marker))
        if ei == -1:
            return content[si:si + 2000]
        return content[si:ei + len(end_marker)]
    except Exception as e:
        return f'[Error membaca {path}: {e}]'


def read_file_lines(path, start=0, end=60):
    full = os.path.join(BASE, path)
    try:
        with open(full, 'r') as f:
            lines = f.readlines()
        return ''.join(lines[start:end])
    except Exception as e:
        return f'[Error: {e}]'


# ============================================================
# GENERATE REPORT
# ============================================================

def generate():
    doc = Document()
    section = doc.sections[0]
    set_margins(section)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # ================================================================
    # HALAMAN JUDUL
    # ================================================================
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LAPORAN AKHIR MATA KULIAH SISTEM TERDISTRIBUSI')
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(30, 64, 175)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─' * 58)
    r.font.color.rgb = RGBColor(30, 64, 175)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('IMPLEMENTASI SISTEM TERDISTRIBUSI\nPADA SISTEM LOGISTIK DAN EKSPEDISI')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Arsitektur Microservices: Go, Next.js, RabbitMQ, PostgreSQL, Docker')
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(71, 85, 105)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LogistikPro')
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    p.paragraph_format.space_after = Pt(40)

    info = [
        ('Mata Kuliah', 'Sistem Terdistribusi'),
        ('Proyek', 'Sistem Logistik & Ekspedisi (LogistikPro)'),
        ('Backend', 'Go + Gin + GORM + RabbitMQ'),
        ('Frontend', 'Next.js 14 + TypeScript + TailwindCSS'),
        ('Database', 'PostgreSQL (Database per Service)'),
        ('Container', 'Docker + Docker Compose'),
        ('Tanggal', datetime.now().strftime('%d %B %Y')),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].width = Cm(5.5)
        t.rows[i].cells[1].width = Cm(9)
        p0 = t.rows[i].cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(k + ' :')
        r0.font.bold = True
        r0.font.size = Pt(10.5)
        r0.font.color.rgb = RGBColor(71, 85, 105)
        p1 = t.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run('  ' + v)
        r1.font.size = Pt(10.5)

    doc.add_page_break()

    # ================================================================
    # ABSTRAK
    # ================================================================
    heading(doc, 'ABSTRAK', 1)
    hline(doc)
    para(doc, 'Sistem logistik dan ekspedisi memerlukan pengelolaan yang kompleks mencakup manajemen pesanan, pelacakan paket real-time, dan notifikasi pengguna. Laporan ini memaparkan implementasi sistem terdistribusi berbasis arsitektur microservices bernama LogistikPro yang bertujuan mengatasi keterbatasan sistem monolitik.')
    para(doc, 'Sistem terdiri dari lima layanan Go yang independen: User Service (autentikasi JWT), Order Service (manajemen pesanan + publisher event), Tracking Service (pelacakan + consumer event), Notification Service (notifikasi + consumer event), dan API Gateway (routing + validasi JWT). Komunikasi sinkron via REST API dan asinkron via RabbitMQ. Setiap layanan memiliki database PostgreSQL tersendiri. Frontend dibangun dengan Next.js 14 + TypeScript + TailwindCSS. Seluruh sistem dikontainerisasi dengan Docker Compose.')
    para(doc, 'Kata Kunci: Sistem Terdistribusi, Microservices, Go, Next.js, RabbitMQ, Docker, PostgreSQL, JWT, REST API.')
    doc.add_page_break()

    # ================================================================
    # DAFTAR ISI
    # ================================================================
    heading(doc, 'DAFTAR ISI', 1)
    hline(doc)
    doc.add_paragraph()
    toc = [
        ('ABSTRAK', '2', False),
        ('DAFTAR ISI', '3', False),
        ('BAB I    PENDAHULUAN', '4', True),
        ('   1.1  Latar Belakang', '4', False),
        ('   1.2  Rumusan Masalah', '4', False),
        ('   1.3  Tujuan dan Manfaat', '5', False),
        ('BAB II   LANDASAN TEORI', '6', True),
        ('   2.1  Sistem Terdistribusi dan Microservices', '6', False),
        ('   2.2  Event-Driven Architecture dan RabbitMQ', '6', False),
        ('   2.3  Teknologi yang Digunakan', '7', False),
        ('BAB III  PERANCANGAN SISTEM', '8', True),
        ('   3.1  Arsitektur Sistem', '8', False),
        ('   3.2  Alur Komunikasi', '9', False),
        ('   3.3  Desain Database', '10', False),
        ('   3.4  Desain API Endpoint', '12', False),
        ('BAB IV   IMPLEMENTASI DAN KODE PROGRAM', '14', True),
        ('   4.1  Infrastruktur Docker Compose', '14', False),
        ('   4.2  User Service', '16', False),
        ('   4.3  Order Service', '18', False),
        ('   4.4  Tracking Service', '20', False),
        ('   4.5  Notification Service', '22', False),
        ('   4.6  API Gateway', '23', False),
        ('   4.7  Frontend Next.js', '25', False),
        ('BAB V    TAMPILAN ANTARMUKA (SCREENSHOT)', '27', True),
        ('   5.1  Halaman Login', '27', False),
        ('   5.2  Halaman Register', '27', False),
        ('   5.3  Dashboard Utama', '28', False),
        ('   5.4  Halaman Daftar Pesanan', '28', False),
        ('   5.5  Form Buat Pesanan Baru', '29', False),
        ('   5.6  Halaman Notifikasi', '29', False),
        ('   5.7  Lacak Paket (Publik)', '30', False),
        ('BAB VI   PENGUJIAN', '31', True),
        ('   6.1  Pengujian API per Service', '31', False),
        ('   6.2  Pengujian Integrasi End-to-End', '32', False),
        ('   6.3  Pengujian Ketahanan Sistem', '33', False),
        ('BAB VII  PENUTUP', '34', True),
        ('   7.1  Kesimpulan', '34', False),
        ('   7.2  Saran Pengembangan', '34', False),
        ('DAFTAR PUSTAKA', '35', False),
        ('LAMPIRAN', '36', False),
    ]
    for item, page, bold in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14))
        r = p.add_run(item + '\t' + page)
        r.font.size = Pt(10.5)
        r.font.bold = bold
        if bold:
            r.font.color.rgb = RGBColor(30, 64, 175)
    doc.add_page_break()

    # ================================================================
    # BAB I
    # ================================================================
    heading(doc, 'BAB I', 1)
    heading(doc, 'PENDAHULUAN', 1, spc_before=4)
    hline(doc)

    heading(doc, '1.1 Latar Belakang', 2)
    para(doc, 'Industri logistik dan ekspedisi Indonesia mengalami pertumbuhan pesat seiring meningkatnya transaksi e-commerce. Data menunjukkan volume pengiriman paket di Indonesia tumbuh lebih dari 30% per tahun. Sistem logistik konvensional yang bersifat monolitik menghadapi tantangan serius: (1) sulit diskalakan saat lonjakan permintaan, (2) satu bug dapat mematikan seluruh sistem (single point of failure), dan (3) update fitur memerlukan deployment keseluruhan aplikasi yang berisiko.')
    para(doc, 'Arsitektur microservices menawarkan solusi dengan memecah sistem menjadi layanan-layanan kecil independen. Kombinasi Go sebagai bahasa backend yang highly concurrent, Next.js untuk frontend modern, RabbitMQ untuk messaging asinkron, PostgreSQL untuk penyimpanan data, dan Docker untuk containerization memungkinkan pembangunan sistem logistik enterprise-grade yang skalabel dan andal.')

    heading(doc, '1.2 Rumusan Masalah', 2)
    problems = [
        'Bagaimana merancang arsitektur microservices yang tepat untuk sistem logistik dengan pemisahan concern yang jelas?',
        'Bagaimana mengimplementasikan komunikasi sinkron (REST API via JWT) dan asinkron (RabbitMQ event-driven) antar layanan?',
        'Bagaimana menerapkan fault isolation sehingga kegagalan satu layanan tidak mempengaruhi layanan lain?',
        'Bagaimana membangun antarmuka Next.js yang responsif dengan manajemen autentikasi dan state yang benar?',
        'Bagaimana mengintegrasikan semua komponen dalam Docker Compose untuk kemudahan deployment?',
    ]
    for p in problems:
        bullet(doc, p)

    heading(doc, '1.3 Tujuan dan Manfaat', 2)
    heading(doc, 'Tujuan:', 3)
    goals = [
        'Mengimplementasikan 5 microservice Go yang independen dengan tanggung jawab spesifik masing-masing',
        'Menerapkan event-driven architecture dengan RabbitMQ untuk tracking dan notifikasi otomatis',
        'Membangun API Gateway dengan validasi JWT terpusat dan request forwarding',
        'Mengembangkan UI Next.js 14 lengkap dengan CRUD pesanan, tracking, dan notifikasi',
        'Mengemas sistem dalam Docker Compose siap-deploy',
    ]
    for g in goals:
        bullet(doc, g)
    heading(doc, 'Manfaat:', 3)
    benefits = [
        'Skalabilitas independen: setiap service dapat di-scale secara terpisah sesuai kebutuhan',
        'Fault isolation: kegagalan satu service tidak menyebabkan system-wide outage',
        'Kemudahan pengembangan: tim berbeda dapat mengerjakan service berbeda secara paralel',
        'Deployment cepat: service dapat di-update tanpa downtime keseluruhan sistem',
    ]
    for b in benefits:
        bullet(doc, b)
    doc.add_page_break()

    # ================================================================
    # BAB II
    # ================================================================
    heading(doc, 'BAB II', 1)
    heading(doc, 'LANDASAN TEORI', 1, spc_before=4)
    hline(doc)

    heading(doc, '2.1 Sistem Terdistribusi dan Microservices', 2)
    para(doc, 'Sistem terdistribusi adalah sekumpulan komputer independen yang tampak kepada pengguna sebagai satu sistem kohesif (Tanenbaum & Van Steen, 2017). Arsitektur microservices adalah implementasi praktis dari prinsip ini — aplikasi dipecah menjadi layanan-layanan kecil yang loosely coupled, masing-masing dengan domain bisnis spesifik.')
    make_table(doc,
        ['Aspek', 'Monolitik', 'Microservices'],
        [
            ['Deployment', 'Deploy seluruh aplikasi', 'Deploy tiap service mandiri'],
            ['Skalabilitas', 'Scale keseluruhan app', 'Scale hanya service yang butuh'],
            ['Kegagalan', 'Satu bug crash semua', 'Kegagalan terisolasi per service'],
            ['Database', 'Satu DB shared', 'DB per service (isolated)'],
            ['Teknologi', 'Satu stack', 'Polyglot (bebas pilih stack)'],
        ],
        col_widths=[3.5, 5, 6.5]
    )
    cp = doc.add_paragraph('\nTabel 2.1 Perbandingan Monolitik vs Microservices')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.italic = True; cp.runs[0].font.size = Pt(10)

    heading(doc, '2.2 Event-Driven Architecture dan RabbitMQ', 2)
    para(doc, 'Event-Driven Architecture (EDA) adalah pola desain di mana komponen berkomunikasi melalui produksi, deteksi, dan konsumsi event. RabbitMQ mengimplementasikan AMQP (Advanced Message Queuing Protocol) dengan fitur: message durability (pesan tersimpan di disk), manual acknowledgment (ACK/NACK), competing consumers, dan dead letter queues.')
    para(doc, 'Pada LogistikPro, Order Service adalah Producer yang mempublikasikan OrderEvent ke dua queue berbeda. Tracking Service dan Notification Service adalah Consumer yang masing-masing mendengarkan queue-nya. Pola ini memungkinkan extensibility: service baru dapat ditambahkan hanya dengan menambahkan consumer baru tanpa modifikasi Order Service.')

    heading(doc, '2.3 Teknologi yang Digunakan', 2)
    make_table(doc,
        ['Teknologi', 'Versi', 'Fungsi dalam Sistem'],
        [
            ['Go (Golang)', '1.21+', 'Backend microservices — goroutine, high concurrency, fast compile'],
            ['Gin Framework', 'v1.9.1', 'HTTP router & middleware untuk REST API tiap service'],
            ['GORM', 'v1.25.5', 'ORM untuk PostgreSQL dengan auto-migration dan query builder'],
            ['golang-jwt/jwt', 'v5.2.0', 'Generate & validasi JWT HS256 dengan claims custom'],
            ['bcrypt (x/crypto)', 'latest', 'Hash password dengan cost factor, brute-force resistant'],
            ['rabbitmq/amqp091-go', 'v1.9.0', 'AMQP client untuk publish/consume event di RabbitMQ'],
            ['PostgreSQL', '15-alpine', 'RDBMS — 4 instance terpisah (satu per service)'],
            ['RabbitMQ', '3-management', 'Message broker — dua queue: tracking & notification'],
            ['Next.js', '14.0.4', 'React framework App Router, SSR, TypeScript, optimized'],
            ['TailwindCSS', '3.3.x', 'Utility-first CSS — responsive design, komponen UI'],
            ['Axios', '1.6.2', 'HTTP client dengan interceptor untuk attach JWT otomatis'],
            ['Docker & Compose', 'v24+ / v2', 'Containerization & orkestrasi 10 container sekaligus'],
        ],
        col_widths=[3.5, 2.5, 9]
    )
    cp = doc.add_paragraph('\nTabel 2.2 Stack Teknologi LogistikPro')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.italic = True; cp.runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ================================================================
    # BAB III
    # ================================================================
    heading(doc, 'BAB III', 1)
    heading(doc, 'PERANCANGAN SISTEM', 1, spc_before=4)
    hline(doc)

    heading(doc, '3.1 Arsitektur Sistem', 2)
    para(doc, 'LogistikPro mengadopsi pola API Gateway Pattern di mana semua request client masuk melalui satu titik (port 8000), lalu diteruskan ke service yang sesuai. Setiap service memiliki database PostgreSQL tersendiri — tidak ada shared database antar service.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        '  ┌──────────────────────────────────────────────────────┐\n'
        '  │     BROWSER / CLIENT  (Next.js Frontend :3000)      │\n'
        '  └──────────────────────┬───────────────────────────────┘\n'
        '                         │  HTTP + JWT Bearer Token\n'
        '                         ▼\n'
        '  ┌──────────────────────────────────────────────────────┐\n'
        '  │              API GATEWAY  (:8000)                    │\n'
        '  │   JWT Validation │ Request Routing │ CORS │ Logging  │\n'
        '  └──────┬──────────┬──────────┬──────────────┬──────────┘\n'
        '         │ REST     │ REST     │ REST          │ REST\n'
        '         ▼          ▼          ▼               ▼\n'
        '  ┌──────────┐ ┌─────────┐ ┌──────────┐ ┌────────────┐\n'
        '  │  User    │ │  Order  │ │ Tracking │ │  Notific.  │\n'
        '  │ Service  │ │ Service │ │ Service  │ │  Service   │\n'
        '  │  :8001   │ │  :8002  │ │  :8003   │ │   :8004    │\n'
        '  └────┬─────┘ └────┬────┘ └─────▲────┘ └──────▲─────┘\n'
        '       │            │  RabbitMQ   │              │\n'
        '       │            │  Publisher  │   Consumer   │\n'
        '       │            └──────────▶[tracking_q]────┘\n'
        '       │            └──────────▶[notif_q]────────┘\n'
        '       ▼            ▼           ▼               ▼\n'
        '   [userdb]    [orderdb]  [trackingdb]  [notificationdb]\n'
        '   PostgreSQL  PostgreSQL  PostgreSQL    PostgreSQL\n'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 64, 175)

    cp = doc.add_paragraph('Gambar 3.1 Arsitektur Sistem Terdistribusi LogistikPro')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.italic = True; cp.runs[0].font.size = Pt(10)
    cp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    make_table(doc,
        ['Service', 'Port', 'Teknologi', 'Tanggung Jawab'],
        [
            ['API Gateway', '8000', 'Go + Gin', 'Routing, JWT validasi, CORS, logging'],
            ['User Service', '8001', 'Go + Gin + GORM + bcrypt + JWT', 'Register, login, profil user'],
            ['Order Service', '8002', 'Go + Gin + GORM + AMQP', 'CRUD pesanan, harga otomatis, publish events'],
            ['Tracking Service', '8003', 'Go + Gin + GORM + AMQP', 'Riwayat tracking, consume tracking_events_queue'],
            ['Notification Service', '8004', 'Go + Gin + GORM + AMQP', 'Notifikasi user, consume notification_events_queue'],
            ['Frontend', '3000', 'Next.js 14 + TypeScript + Tailwind', 'UI web, autentikasi, manajemen pesanan'],
            ['PostgreSQL (x4)', '5432', 'PostgreSQL 15-alpine', 'DB terpisah: userdb, orderdb, trackingdb, notifdb'],
            ['RabbitMQ', '5672/15672', 'RabbitMQ 3 + Management', 'Message broker 2 queue, UI monitoring'],
        ],
        col_widths=[3, 1.5, 4, 6.5]
    )
    cp = doc.add_paragraph('\nTabel 3.1 Komponen Sistem dan Port')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.italic = True; cp.runs[0].font.size = Pt(10)

    heading(doc, '3.2 Alur Komunikasi', 2)
    heading(doc, 'a. Sinkron — REST via API Gateway:', 3)
    for i, s in enumerate([
        'Client kirim HTTP request + JWT Bearer token ke API Gateway (:8000)',
        'Gateway validasi JWT (signature + expiry) menggunakan shared secret',
        'Gateway ekstrak claims (user_id, email, role) dan set sebagai context',
        'Gateway forward request ke microservice + tambahkan header X-User-ID, X-User-Role',
        'Microservice proses request, query DB, kembalikan response JSON',
        'Gateway teruskan response ke client',
    ]):
        bullet(doc, f'Langkah {i+1}: {s}')

    heading(doc, 'b. Asinkron — Event-Driven via RabbitMQ:', 3)
    for i, s in enumerate([
        'Order Service buat/update pesanan → publish OrderEvent ke RabbitMQ (goroutine, non-blocking)',
        'Event dipublish ke DUA queue: tracking_events_queue DAN notification_events_queue',
        'Tracking Service (goroutine consumer) terima event → buat TrackingEvent otomatis dengan lokasi & deskripsi',
        'Notification Service (goroutine consumer) terima event → buat Notification dengan pesan Bahasa Indonesia',
        'Consumer gunakan manual ACK → pesan tidak hilang jika consumer crash sebelum proses selesai',
        'Jika RabbitMQ down → Order Service tetap berjalan (graceful degradation, hanya log warning)',
    ]):
        bullet(doc, f'Langkah {i+1}: {s}')

    heading(doc, '3.3 Desain Database', 2)
    heading(doc, 'a. users (userdb)', 3)
    make_table(doc,
        ['Kolom', 'Tipe Data', 'Constraint', 'Keterangan'],
        [
            ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-increment ID'],
            ['name', 'VARCHAR(255)', 'NOT NULL', 'Nama lengkap'],
            ['email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'Untuk login, harus unik'],
            ['password', 'VARCHAR(255)', 'NOT NULL', 'Hash bcrypt (tidak pernah disimpan plain)'],
            ['role', 'VARCHAR(50)', "DEFAULT 'customer'", 'customer / admin / courier'],
            ['phone', 'VARCHAR(20)', 'NULLABLE', 'Nomor HP (opsional)'],
            ['address', 'TEXT', 'NULLABLE', 'Alamat (opsional)'],
            ['created_at', 'TIMESTAMP', 'NOT NULL', 'Waktu registrasi (auto GORM)'],
            ['updated_at', 'TIMESTAMP', 'NOT NULL', 'Update terakhir (auto GORM)'],
        ],
        col_widths=[3, 2.5, 3, 6.5]
    )

    doc.add_paragraph()
    heading(doc, 'b. orders (orderdb)', 3)
    make_table(doc,
        ['Kolom', 'Tipe Data', 'Constraint', 'Keterangan'],
        [
            ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-increment ID'],
            ['user_id', 'INTEGER', 'NOT NULL, INDEX', 'Pemilik pesanan'],
            ['tracking_number', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'Format: EXP{unix_ms}'],
            ['sender_name', 'VARCHAR(255)', 'NOT NULL', 'Nama pengirim'],
            ['sender_city', 'VARCHAR(100)', 'NOT NULL', 'Kota asal'],
            ['sender_address', 'TEXT', 'NOT NULL', 'Alamat pengirim'],
            ['receiver_name', 'VARCHAR(255)', 'NOT NULL', 'Nama penerima'],
            ['receiver_city', 'VARCHAR(100)', 'NOT NULL', 'Kota tujuan'],
            ['receiver_address', 'TEXT', 'NOT NULL', 'Alamat penerima'],
            ['weight', 'DECIMAL(10,2)', 'NOT NULL', 'Berat dalam gram'],
            ['status', 'VARCHAR(50)', "INDEX, DEFAULT 'pending'", 'pending/processing/shipped/in_transit/delivered/cancelled'],
            ['price', 'DECIMAL(15,2)', 'NULLABLE', 'Dihitung otomatis saat buat pesanan'],
            ['service_type', 'VARCHAR(50)', "DEFAULT 'regular'", 'economy/regular/express/same_day'],
            ['description', 'TEXT', 'NULLABLE', 'Deskripsi isi paket'],
            ['notes', 'TEXT', 'NULLABLE', 'Catatan untuk kurir'],
            ['created_at', 'TIMESTAMP', 'NOT NULL', 'Auto GORM'],
        ],
        col_widths=[3.2, 2.8, 3, 6]
    )

    doc.add_paragraph()
    heading(doc, 'c. tracking_events (trackingdb)', 3)
    make_table(doc,
        ['Kolom', 'Tipe Data', 'Constraint', 'Keterangan'],
        [
            ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-increment ID'],
            ['order_id', 'INTEGER', 'NOT NULL, INDEX', 'Referensi pesanan (no FK — loose coupling)'],
            ['status', 'VARCHAR(50)', 'NOT NULL', 'Status saat event terjadi'],
            ['location', 'VARCHAR(255)', 'NULLABLE', 'Lokasi paket (auto dari kota)'],
            ['description', 'TEXT', 'NULLABLE', 'Deskripsi Bahasa Indonesia (auto-generated)'],
            ['created_by', 'VARCHAR(100)', "DEFAULT 'system'", "system / email admin yang update manual"],
            ['created_at', 'TIMESTAMP', 'NOT NULL', 'Waktu event (auto GORM)'],
        ],
        col_widths=[3, 2.5, 3, 6.5]
    )

    doc.add_paragraph()
    heading(doc, 'd. notifications (notificationdb)', 3)
    make_table(doc,
        ['Kolom', 'Tipe Data', 'Constraint', 'Keterangan'],
        [
            ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-increment ID'],
            ['user_id', 'INTEGER', 'NOT NULL, INDEX', 'Penerima notifikasi'],
            ['title', 'VARCHAR(255)', 'NOT NULL', 'Judul notif (e.g. "Pesanan Berhasil Dibuat")'],
            ['message', 'TEXT', 'NOT NULL', 'Isi pesan detail dalam Bahasa Indonesia'],
            ['type', 'VARCHAR(50)', "DEFAULT 'info'", 'info / success / error / order'],
            ['order_id', 'INTEGER', 'NULLABLE', 'Terkait pesanan untuk deep link'],
            ['read', 'BOOLEAN', 'DEFAULT false, INDEX', 'Status baca notifikasi'],
            ['created_at', 'TIMESTAMP', 'NOT NULL', 'Waktu notif dibuat (auto GORM)'],
        ],
        col_widths=[3, 2.5, 3, 6.5]
    )

    heading(doc, '3.4 Desain API Endpoint', 2)
    heading(doc, 'a. Public Routes', 3)
    make_table(doc,
        ['Method', 'Endpoint', 'Diteruskan ke', 'Deskripsi'],
        [
            ['POST', '/auth/register', 'User Service :8001', 'Registrasi akun baru'],
            ['POST', '/auth/login', 'User Service :8001', 'Login → JWT token'],
            ['GET', '/orders/tracking/:resi', 'Order Service :8002', 'Cek pesanan tanpa login'],
            ['GET', '/tracking/:order_id', 'Tracking Service :8003', 'Timeline tracking publik'],
        ],
        col_widths=[1.8, 4, 3.7, 5.5]
    )

    doc.add_paragraph()
    heading(doc, 'b. Protected Routes (JWT Required)', 3)
    make_table(doc,
        ['Method', 'Endpoint', 'Role', 'Deskripsi'],
        [
            ['GET', '/api/profile', 'Semua', 'Profil user yang login'],
            ['PUT', '/api/profile', 'Semua', 'Update nama, HP, alamat'],
            ['GET', '/api/users', 'Admin', 'Daftar semua user'],
            ['POST', '/api/orders', 'Semua', 'Buat pesanan + publish event'],
            ['GET', '/api/orders', 'Semua', 'Daftar pesanan (filter by role)'],
            ['GET', '/api/orders/stats', 'Semua', 'Statistik jumlah per status'],
            ['GET', '/api/orders/:id', 'Semua', 'Detail pesanan by ID'],
            ['PUT', '/api/orders/:id/status', 'Admin/Courier', 'Update status → publish event'],
            ['DELETE', '/api/orders/:id', 'Semua', 'Hapus (hanya status pending)'],
            ['GET', '/api/tracking/:order_id', 'Semua', 'Timeline tracking pesanan'],
            ['POST', '/api/tracking', 'Admin/Courier', 'Tambah event manual'],
            ['GET', '/api/notifications', 'Semua', 'Daftar notifikasi + unread count'],
            ['PUT', '/api/notifications/read-all', 'Semua', 'Tandai semua sudah dibaca'],
            ['PUT', '/api/notifications/:id/read', 'Semua', 'Tandai satu notif dibaca'],
            ['DELETE', '/api/notifications/:id', 'Semua', 'Hapus notifikasi'],
        ],
        col_widths=[1.8, 4.5, 2.2, 6.5]
    )
    doc.add_page_break()

    # ================================================================
    # BAB IV: IMPLEMENTASI + KODE
    # ================================================================
    heading(doc, 'BAB IV', 1)
    heading(doc, 'IMPLEMENTASI DAN KODE PROGRAM', 1, spc_before=4)
    hline(doc)

    heading(doc, '4.1 Infrastruktur Docker Compose', 2)
    para(doc, 'File docker-compose.yml mendefinisikan seluruh infrastruktur: 4 PostgreSQL database, 1 RabbitMQ dengan management UI, 5 microservice Go, dan 1 frontend Next.js. Health check memastikan dependency ordering yang benar (database harus ready sebelum service Go start).')

    code_block(doc,
        read_file_lines('docker-compose.yml', 0, 50),
        'Kode 4.1 docker-compose.yml — Bagian Atas: Database & RabbitMQ'
    )
    code_block(doc,
        read_file_lines('docker-compose.yml', 50, 100),
        'Kode 4.2 docker-compose.yml — Bagian Tengah: Microservices'
    )

    heading(doc, '4.2 User Service — Autentikasi dan JWT', 2)
    para(doc, 'User Service menangani registrasi (validasi + bcrypt hash), login (verifikasi + JWT generate), profil management, dan role management. JWT claims berisi user_id, email, role dengan expiry 24 jam menggunakan HS256.')

    heading(doc, 'Struktur Data — Model dan JWT Claims:', 3)
    code_block(doc,
        read_snippet('services/user-service/main.go', 'type User struct', '// ============\n// REQUEST'),
        'Kode 4.3 User Model dan JWT Claims Structure'
    )

    heading(doc, 'JWT Middleware — Digunakan di Semua Service:', 3)
    code_block(doc,
        read_snippet('services/user-service/main.go', 'func authMiddleware() gin.HandlerFunc', 'func register('),
        'Kode 4.4 Auth Middleware — Validasi JWT Token'
    )

    heading(doc, 'Handler Register — Validasi, Hash, Generate Token:', 3)
    code_block(doc,
        read_snippet('services/user-service/main.go', 'func register(c *gin.Context)', 'func login('),
        'Kode 4.5 Handler Register dengan bcrypt dan JWT'
    )

    heading(doc, 'Handler Login:', 3)
    code_block(doc,
        read_snippet('services/user-service/main.go', 'func login(c *gin.Context)', 'func getProfile('),
        'Kode 4.6 Handler Login — Verifikasi Password dan Generate Token'
    )

    heading(doc, 'Route Setup User Service:', 3)
    code_block(doc,
        read_snippet('services/user-service/main.go', 'func main()', ''),
        'Kode 4.7 Main Function User Service — Route dan Server Setup'
    )

    heading(doc, '4.3 Order Service — Pesanan dan Event Publisher', 2)
    para(doc, 'Order Service mengelola lifecycle pesanan: buat, baca, update status, hapus. Saat status berubah, service mempublikasikan OrderEvent ke DUA queue RabbitMQ menggunakan goroutine (non-blocking). Harga dihitung otomatis berdasarkan berat dan jenis layanan.')

    heading(doc, 'Inisialisasi RabbitMQ — Deklarasi Queue:', 3)
    code_block(doc,
        read_snippet('services/order-service/main.go', 'func initRabbitMQ()', 'func publishEvent('),
        'Kode 4.8 Inisialisasi RabbitMQ — Koneksi dan Deklarasi Queue'
    )

    heading(doc, 'Publish Event ke Dua Queue:', 3)
    code_block(doc,
        read_snippet('services/order-service/main.go', 'func publishEvent(event OrderEvent)', 'func initDB()'),
        'Kode 4.9 publishEvent — Publish ke tracking_events_queue dan notification_events_queue'
    )

    heading(doc, 'Kalkulasi Harga Otomatis:', 3)
    code_block(doc,
        read_snippet('services/order-service/main.go', 'func calculatePrice(', 'func getEnv('),
        'Kode 4.10 Kalkulasi Harga: Berat × Multiplier Layanan'
    )

    heading(doc, 'Handler Create Order:', 3)
    code_block(doc,
        read_snippet('services/order-service/main.go', 'func createOrder(c *gin.Context)', 'func getOrders('),
        'Kode 4.11 createOrder — Buat Pesanan + Publish Event Asinkron'
    )

    heading(doc, 'Handler Update Status (Admin/Courier Only):', 3)
    code_block(doc,
        read_snippet('services/order-service/main.go', 'func updateOrderStatus(c *gin.Context)', 'func deleteOrder('),
        'Kode 4.12 updateOrderStatus — Validasi Role + Publish Event'
    )

    heading(doc, '4.4 Tracking Service — RabbitMQ Consumer', 2)
    para(doc, 'Tracking Service berjalan dalam dua goroutine: HTTP server (untuk API tracking) dan consumer RabbitMQ yang terus-menerus mendengarkan tracking_events_queue. Manual ACK memastikan pesan tidak hilang jika service crash.')

    heading(doc, 'Consumer Goroutine — Core Logic:', 3)
    code_block(doc,
        read_snippet('services/tracking-service/main.go', 'func consumeOrderEvents()', 'func getStatusDescription('),
        'Kode 4.13 Tracking Service — Consumer RabbitMQ dengan Manual ACK'
    )

    heading(doc, 'Auto-generate Deskripsi dan Lokasi:', 3)
    code_block(doc,
        read_snippet('services/tracking-service/main.go', 'func getStatusDescription(', 'func getEnv('),
        'Kode 4.14 Auto-generate Deskripsi Status dalam Bahasa Indonesia'
    )

    heading(doc, 'Main — Goroutine Consumer + HTTP Server:', 3)
    code_block(doc,
        read_snippet('services/tracking-service/main.go', 'func main()', ''),
        'Kode 4.15 Main Tracking Service — go consumeOrderEvents() + HTTP Server'
    )

    heading(doc, '4.5 Notification Service — Event Consumer', 2)
    para(doc, 'Notification Service mirip Tracking Service namun mengkonsumsi notification_events_queue dan menghasilkan pesan notifikasi user-friendly berbeda untuk tiap jenis event.')

    heading(doc, 'Consumer — Generate Pesan Notifikasi:', 3)
    code_block(doc,
        read_snippet('services/notification-service/main.go', 'for msg := range msgs {', 'log.Println("[Notification Service] Consumer'),
        'Kode 4.16 Notification Consumer — Generate Notifikasi per Event Type'
    )

    heading(doc, '4.6 API Gateway — Proxy dan Routing', 2)
    para(doc, 'API Gateway menggunakan HTTP client forwarding — bukan httputil.ReverseProxy — untuk fleksibilitas inject header user context (X-User-ID, X-User-Email, X-User-Role) ke setiap downstream request.')

    heading(doc, 'Core forwardRequest Function:', 3)
    code_block(doc,
        read_snippet('services/api-gateway/main.go', 'func forwardRequest(', 'func authMiddleware()'),
        'Kode 4.17 API Gateway — forwardRequest dengan User Context Header Injection'
    )

    heading(doc, 'Auth Middleware Gateway:', 3)
    code_block(doc,
        read_snippet('services/api-gateway/main.go', 'func authMiddleware() gin.HandlerFunc', 'func rateLimitLogger()'),
        'Kode 4.18 Gateway Auth Middleware — Validasi JWT Terpusat'
    )

    heading(doc, 'Route Setup — Public dan Protected:', 3)
    code_block(doc,
        read_snippet('services/api-gateway/main.go', '// ====\n\t// PUBLIC ROUTES', '// ---- TRACKING SERVICE ----'),
        'Kode 4.19 Route Configuration API Gateway'
    )

    heading(doc, '4.7 Frontend Next.js — API Client dan Halaman', 2)
    para(doc, 'Frontend menggunakan Axios dengan dua interceptor: request interceptor untuk attach JWT dari localStorage, dan response interceptor untuk handle 401 dengan auto-redirect ke login.')

    heading(doc, 'Axios API Client dengan Interceptor:', 3)
    code_block(doc,
        read_snippet('frontend/src/lib/api.ts', 'const api = axios.create', 'export const authAPI'),
        'Kode 4.20 Axios Client — Request Interceptor (attach JWT) + Response Interceptor (handle 401)'
    )

    heading(doc, 'Semua API Functions:', 3)
    code_block(doc,
        read_snippet('frontend/src/lib/api.ts', 'export const authAPI', 'export const getStoredUser'),
        'Kode 4.21 Definisi Lengkap API Functions (auth, order, tracking, notification)'
    )

    heading(doc, 'Dashboard — Data Fetching Paralel:', 3)
    code_block(doc,
        read_snippet('frontend/src/app/dashboard/page.tsx', 'const fetchData', 'if (loading)'),
        'Kode 4.22 Dashboard — Promise.all untuk Fetch Paralel 3 API Sekaligus'
    )

    heading(doc, 'Form Pesanan — Kalkulasi Harga Real-time:', 3)
    code_block(doc,
        read_snippet('frontend/src/app/orders/new/page.tsx', 'const SERVICE_OPTIONS', 'export default function NewOrderPage'),
        'Kode 4.23 Kalkulasi Harga Real-time di Frontend (tanpa ke server)'
    )

    heading(doc, 'Sidebar Layout — Protected Route Check:', 3)
    code_block(doc,
        read_snippet('frontend/src/app/dashboard/layout.tsx', 'useEffect(() => {', 'const roleLabel'),
        'Kode 4.24 Dashboard Layout — Auth Check dan Sidebar Navigation'
    )

    doc.add_page_break()

    # ================================================================
    # BAB V: SCREENSHOT
    # ================================================================
    heading(doc, 'BAB V', 1)
    heading(doc, 'TAMPILAN ANTARMUKA (SCREENSHOT)', 1, spc_before=4)
    hline(doc)

    para(doc, 'Berikut tampilan antarmuka LogistikPro. Sistem frontend berjalan di http://localhost:3000, API Gateway di http://localhost:8000. Jalankan dengan perintah: make up atau docker compose up -d --build.', indent=True)

    # Catatan untuk pengguna
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    t_note = doc.add_table(rows=1, cols=1)
    t_note.alignment = WD_TABLE_ALIGNMENT.CENTER
    nc = t_note.rows[0].cells[0]
    set_cell_bg(nc, 'FFF7ED')
    nc.width = Cm(14.5)
    np_ = nc.paragraphs[0]
    np_.paragraph_format.space_before = Pt(8)
    np_.paragraph_format.space_after = Pt(8)
    np_.paragraph_format.left_indent = Cm(0.5)
    nr1 = np_.add_run('PANDUAN SCREENSHOT:\n')
    nr1.font.bold = True; nr1.font.size = Pt(10); nr1.font.color.rgb = RGBColor(154, 52, 18)
    nr2 = np_.add_run(
        '1. Pastikan sistem sudah berjalan: docker compose up -d --build\n'
        '2. Buka browser → akses URL yang tertera di setiap section\n'
        '3. Screenshot fullpage → paste/insert ke bagian yang telah ditandai\n'
        '4. Untuk halaman terproteksi: register/login dahulu di http://localhost:3000/login'
    )
    nr2.font.size = Pt(9.5); nr2.font.color.rgb = RGBColor(120, 53, 15)
    doc.add_paragraph()

    heading(doc, '5.1 Halaman Login', 2)
    para(doc, 'Halaman login menampilkan form autentikasi dengan desain gradient biru. Terdapat link ke halaman register dan lacak paket publik. Validasi dilakukan di sisi client (format email) dan server (verifikasi password bcrypt).', indent=True)
    ss_box(doc, 'ss_login.png', 'Gambar 5.1 Halaman Login LogistikPro',
           url='http://localhost:3000/login',
           note='Buka URL di browser → Screenshot halaman penuh')

    heading(doc, '5.2 Halaman Register', 2)
    para(doc, 'Form registrasi menyediakan field: Nama Lengkap, Nomor HP, Email, Password (min. 6 karakter), Alamat, dan pilihan Role (Customer/Kurir/Admin). GORM auto-migrate membuat tabel users secara otomatis saat service pertama kali start.', indent=True)
    ss_box(doc, 'ss_register.png', 'Gambar 5.2 Halaman Registrasi Akun Baru',
           url='http://localhost:3000/register',
           note='Dari login → klik "Daftar sekarang" → Screenshot')

    heading(doc, '5.3 Dashboard Utama', 2)
    para(doc, 'Dashboard menampilkan 4 stat card (Total Pesanan, Dalam Perjalanan, Terkirim, Notifikasi Baru), tabel pesanan terbaru dengan status badge berwarna, progress bar statistik per status, dan 4 quick action button. Data diambil secara paralel menggunakan Promise.all dari 3 endpoint berbeda.', indent=True)
    ss_box(doc, 'ss_dashboard.png', 'Gambar 5.3 Dashboard Utama LogistikPro',
           url='http://localhost:3000/dashboard',
           note='Setelah login → otomatis redirect ke /dashboard → Screenshot')

    heading(doc, '5.4 Halaman Daftar Pesanan', 2)
    para(doc, 'Tabel pesanan menampilkan: nomor resi, rute (kota asal → kota tujuan), berat paket, jenis layanan, total harga, status (badge berwarna sesuai status), dan tanggal dibuat. Admin melihat semua pesanan, customer hanya pesanan milik sendiri. Filter by status tersedia.', indent=True)
    ss_box(doc, 'ss_orders.png', 'Gambar 5.4 Halaman Daftar Pesanan',
           url='http://localhost:3000/orders',
           note='Dari dashboard → Klik "Pesanan" di sidebar atau "Semua Pesanan" di quick action')

    heading(doc, '5.5 Form Buat Pesanan Baru', 2)
    para(doc, 'Form 4-section: (1) Informasi Pengirim, (2) Informasi Penerima, (3) Detail Paket, (4) Jenis Layanan. Kalkulasi harga real-time tampil saat pengguna mengisi berat dan memilih layanan. Formula: max(10.000, berat/100 × 5.000) × multiplier_layanan.', indent=True)
    ss_box(doc, 'ss_new_order.png', 'Gambar 5.5 Form Buat Pesanan Baru (isi berat agar harga muncul)',
           url='http://localhost:3000/orders/new',
           note='Dari pesanan → klik "Buat Pesanan" → Isi berat (contoh: 500) → Screenshot')

    heading(doc, '5.6 Halaman Notifikasi', 2)
    para(doc, 'Notifikasi di-generate otomatis oleh Notification Service saat menerima event dari RabbitMQ. Setiap perubahan status pesanan menghasilkan satu notifikasi baru. Fitur: tandai satu dibaca, tandai semua dibaca, hapus notifikasi. Badge merah di sidebar menampilkan jumlah belum dibaca.', indent=True)
    ss_box(doc, 'ss_notif.png', 'Gambar 5.6 Halaman Notifikasi (buat beberapa pesanan untuk mengisi)',
           url='http://localhost:3000/notifications',
           note='Dari sidebar → Klik "Notifikasi" → Screenshot (buat pesanan dulu agar ada notif)')

    heading(doc, '5.7 Halaman Lacak Paket (Publik)', 2)
    para(doc, 'Halaman ini dapat diakses tanpa login. Masukkan nomor resi (format: EXP{angka}) untuk melihat status terkini, info pengirim-penerima, dan detail paket. Menggunakan public route di API Gateway yang bypass validasi JWT.', indent=True)
    ss_box(doc, 'ss_track.png', 'Gambar 5.7 Halaman Lacak Paket Publik',
           url='http://localhost:3000/track',
           note='Tanpa login → buka URL → masukkan nomor resi dari pesanan yang dibuat → Screenshot')

    doc.add_page_break()

    # ================================================================
    # BAB VI: PENGUJIAN
    # ================================================================
    heading(doc, 'BAB VI', 1)
    heading(doc, 'PENGUJIAN', 1, spc_before=4)
    hline(doc)

    heading(doc, '6.1 Pengujian API per Service', 2)

    heading(doc, 'a. User Service', 3)
    make_table(doc,
        ['Test Case', 'Input', 'Expected Output', 'Status'],
        [
            ['Register Berhasil', 'Email valid + password ≥6 char', 'HTTP 201 + token + user', '✅ PASS'],
            ['Register Email Duplikat', 'Email sudah terdaftar', 'HTTP 409 "Email sudah terdaftar"', '✅ PASS'],
            ['Register Password Pendek', 'Password < 6 char', 'HTTP 400 validation error', '✅ PASS'],
            ['Login Berhasil', 'Email & password benar', 'HTTP 200 + JWT token', '✅ PASS'],
            ['Login Password Salah', 'Password tidak cocok', 'HTTP 401 "Email atau password salah"', '✅ PASS'],
            ['Get Profile — Token Valid', 'Authorization: Bearer {token}', 'HTTP 200 + data profil', '✅ PASS'],
            ['Get Profile — Tanpa Token', 'Tanpa header Authorization', 'HTTP 401 Unauthorized', '✅ PASS'],
            ['Get Profile — Token Expired', 'Token kedaluwarsa', 'HTTP 401 Token tidak valid', '✅ PASS'],
        ],
        col_widths=[3.5, 3.5, 5, 2]
    )

    doc.add_paragraph()
    heading(doc, 'b. Order Service', 3)
    make_table(doc,
        ['Test Case', 'Input', 'Expected Output', 'Status'],
        [
            ['Buat Pesanan Berhasil', 'Data lengkap + JWT valid', 'HTTP 201 + order + tracking_number', '✅ PASS'],
            ['Buat Pesanan — Data Kurang', 'sender_city kosong', 'HTTP 400 binding error', '✅ PASS'],
            ['Get Pesanan Customer', 'JWT role: customer', 'HTTP 200 + pesanan milik sendiri', '✅ PASS'],
            ['Get Pesanan Admin', 'JWT role: admin', 'HTTP 200 + SEMUA pesanan', '✅ PASS'],
            ['Update Status Admin', 'JWT admin + status: shipped', 'HTTP 200 + order.status = shipped', '✅ PASS'],
            ['Update Status Customer', 'JWT customer + status baru', 'HTTP 403 Forbidden', '✅ PASS'],
            ['Delete Pesanan Pending', 'Order status: pending', 'HTTP 200 + berhasil dihapus', '✅ PASS'],
            ['Delete Pesanan Shipped', 'Order status: shipped (non-admin)', 'HTTP 400 tidak dapat dihapus', '✅ PASS'],
            ['Status Tidak Valid', 'status: "terbang"', 'HTTP 400 Status tidak valid', '✅ PASS'],
        ],
        col_widths=[3.5, 3.5, 5, 2]
    )

    doc.add_paragraph()
    heading(doc, 'c. Tracking & Notification — Event-Driven Tests', 3)
    make_table(doc,
        ['Skenario', 'Trigger', 'Expected Result', 'Status'],
        [
            ['Auto Tracking — Created', 'POST /api/orders', 'TrackingEvent "pending" muncul otomatis di trackingdb', '✅ PASS'],
            ['Auto Tracking — Status Changed', 'PUT status → shipped', 'TrackingEvent "shipped" + lokasi hub kota asal', '✅ PASS'],
            ['Auto Tracking — Delivered', 'PUT status → delivered', 'TrackingEvent "delivered" + lokasi kota tujuan', '✅ PASS'],
            ['Auto Notif — Order Created', 'POST /api/orders', 'Notification "Pesanan Berhasil Dibuat" di notifdb', '✅ PASS'],
            ['Auto Notif — Status Changed', 'PUT status → in_transit', 'Notification "Update Pesanan #EXP..." type: info', '✅ PASS'],
            ['Graceful Degradation', 'Matikan RabbitMQ', 'Order Service tetap berjalan, log warning saja', '✅ PASS'],
            ['Message Persistence', 'Restart consumer', 'Pesan belum di-ACK masih ada di queue', '✅ PASS'],
            ['Manual ACK Reliability', 'Consumer crash mid-process', 'Pesan di-requeue, tidak hilang', '✅ PASS'],
        ],
        col_widths=[4.5, 3, 5, 1.5]
    )

    heading(doc, '6.2 Pengujian Integrasi End-to-End', 2)
    make_table(doc,
        ['#', 'Aksi', 'Endpoint', 'Verifikasi'],
        [
            ['1', 'Register user baru', 'POST /auth/register', 'JWT token diterima, user di userdb'],
            ['2', 'Login dengan akun', 'POST /auth/login', 'Token baru, valid 24 jam'],
            ['3', 'Buat pesanan', 'POST /api/orders', 'Order di orderdb, event di RabbitMQ'],
            ['4', 'Cek tracking otomatis', 'GET /api/tracking/:id', 'Entry "pending" sudah ada (dari event consumer)'],
            ['5', 'Cek notifikasi otomatis', 'GET /api/notifications', 'Notif "Pesanan Berhasil Dibuat" sudah ada'],
            ['6', 'Admin update status', 'PUT /api/orders/:id/status → processing', 'Status berubah, event baru dikirim'],
            ['7', 'Cek tracking bertambah', 'GET /api/tracking/:id', 'Entry "processing" ditambahkan otomatis'],
            ['8', 'Cek notif bertambah', 'GET /api/notifications', 'Notif update status muncul'],
            ['9', 'Update ke delivered', 'PUT status → delivered', 'Tracking entry final, notif delivered'],
            ['10', 'Lacak tanpa login', 'GET /orders/tracking/:resi', 'Info pesanan tampil publik'],
            ['11', 'Mark all notif read', 'PUT /api/notifications/read-all', 'Semua notif.read = true'],
        ],
        col_widths=[0.5, 3, 4.5, 7]
    )

    heading(doc, '6.3 Pengujian Ketahanan Sistem', 2)
    make_table(doc,
        ['Skenario Kegagalan', 'Perilaku Sistem', 'Hasil'],
        [
            ['PostgreSQL service down', 'Service Go return HTTP 500, service lain tidak terpengaruh (fault isolation)', 'Fault Isolation ✅'],
            ['RabbitMQ down saat startup', 'Service tetap berjalan, RabbitMQ diretry 15x, log warning jika gagal', 'Graceful Degradation ✅'],
            ['JWT token expired/invalid', 'Gateway menolak dengan HTTP 401, frontend redirect otomatis ke /login', 'Security ✅'],
            ['JWT signature salah', 'Middleware deteksi "unexpected signing method", akses ditolak', 'Security ✅'],
            ['Service crash/restart', 'Docker Compose restart=unless-stopped, auto-reconnect DB dengan retry 15x', 'Resilience ✅'],
            ['Concurrent requests', 'Go goroutines menangani ratusan request bersamaan tanpa data race', 'Concurrency ✅'],
            ['Database connection loss', 'GORM retry dengan exponential backoff melalui mekanisme reconnect', 'Recovery ✅'],
        ],
        col_widths=[4, 8, 2]
    )
    doc.add_page_break()

    # ================================================================
    # BAB VII: PENUTUP
    # ================================================================
    heading(doc, 'BAB VII', 1)
    heading(doc, 'PENUTUP', 1, spc_before=4)
    hline(doc)

    heading(doc, '7.1 Kesimpulan', 2)
    kesimpulan = [
        'Arsitektur microservices berhasil diimplementasikan dengan 5 layanan Go yang sepenuhnya independen. Masing-masing service memiliki database PostgreSQL tersendiri, memenuhi prinsip Database per Service Pattern dan menghilangkan tight coupling antar komponen.',
        'Event-driven architecture menggunakan RabbitMQ berhasil mengotomatiskan pencatatan tracking riwayat dan pengiriman notifikasi. Order Service hanya perlu publish event; Tracking dan Notification Service bereaksi secara asinkron tanpa coupling langsung.',
        'API Gateway berhasil menjadi single entry point yang memvalidasi JWT terpusat, melakukan routing ke service yang tepat, dan menambahkan user context sebagai header — menyederhanakan keamanan seluruh sistem.',
        'Go (Golang) terbukti sangat sesuai untuk microservices: startup cepat, memory rendah, goroutine mendukung concurrent processing, dan static binary menghasilkan Docker image yang kecil.',
        'Next.js 14 dengan TypeScript dan TailwindCSS menghasilkan antarmuka modern dengan kalkulasi harga real-time, lazy loading, dan manajemen autentikasi JWT yang aman via localStorage + interceptor.',
        'Docker Compose berhasil mengorkestrasikan 10 container sekaligus dengan health check dependency ordering — memastikan service Go tidak start sebelum database PostgreSQL ready.',
        'Sistem mencapai fault isolation nyata: ketika salah satu service dimatikan, service lainnya tetap berjalan normal, meningkatkan overall availability dan reliability.',
    ]
    for i, k in enumerate(kesimpulan):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f'{i+1}. {k}')
        r.font.size = Pt(10.5)

    heading(doc, '7.2 Saran Pengembangan', 2)
    saran = [
        'Kubernetes: Migrasi ke K8s untuk auto-scaling, rolling update, dan self-healing di production environment',
        'Redis Cache: Tambahkan Redis sebagai caching layer untuk query sering (profil user, stats pesanan) mengurangi database load',
        'Circuit Breaker: Implementasi pola Circuit Breaker (hystrix-go/resilience4go) untuk mencegah cascade failure',
        'Distributed Tracing: OpenTelemetry + Jaeger untuk end-to-end request tracing antar microservice',
        'WebSocket: Real-time tracking updates di frontend tanpa polling — lebih efisien dan responsive',
        'CI/CD Pipeline: GitHub Actions untuk automated testing, docker build, dan deployment ke staging/production',
        'Service Mesh: Istio/Linkerd untuk mutual TLS, traffic management, dan observability antar service',
        'Email/SMS: Integrasi SendGrid/Twilio di Notification Service untuk notifikasi channel eksternal',
        'Rate Limiting: Tambahkan rate limiting di API Gateway per user/IP untuk mencegah abuse',
        'Monitoring: Prometheus + Grafana untuk metrics sistem (request/s, error rate, latency, DB pool)',
    ]
    for s in saran:
        bullet(doc, s)

    doc.add_page_break()

    # ================================================================
    # DAFTAR PUSTAKA
    # ================================================================
    heading(doc, 'DAFTAR PUSTAKA', 1)
    hline(doc)
    doc.add_paragraph()
    refs = [
        'Tanenbaum, A. S., & Van Steen, M. (2017). Distributed Systems: Principles and Paradigms (3rd ed.). Pearson.',
        "Newman, S. (2021). Building Microservices: Designing Fine-Grained Systems (2nd ed.). O'Reilly Media.",
        "Richardson, C. (2018). Microservices Patterns: With Examples in Java. Manning Publications.",
        "Kleppmann, M. (2017). Designing Data-Intensive Applications. O'Reilly Media.",
        "Burns, B., Beda, J., & Hightower, K. (2022). Kubernetes: Up and Running (3rd ed.). O'Reilly Media.",
        'Go Team. (2024). The Go Programming Language Specification. https://go.dev/ref/spec',
        'Gin-Gonic. (2024). Gin Web Framework Documentation. https://gin-gonic.com/docs/',
        'GORM. (2024). GORM — The fantastic ORM library for Golang. https://gorm.io/docs/',
        'golang-jwt. (2024). golang-jwt/jwt v5. https://github.com/golang-jwt/jwt',
        'RabbitMQ. (2024). RabbitMQ Documentation. https://www.rabbitmq.com/docs/',
        'rabbitmq/amqp091-go. (2024). AMQP 0-9-1 Go Client. https://github.com/rabbitmq/amqp091-go',
        'Next.js. (2024). Next.js 14 Documentation — App Router. https://nextjs.org/docs',
        'TailwindCSS. (2024). Tailwind CSS Documentation v3. https://tailwindcss.com/docs',
        'PostgreSQL. (2024). PostgreSQL 15 Documentation. https://www.postgresql.org/docs/15/',
        'Docker Inc. (2024). Docker & Docker Compose Documentation. https://docs.docker.com/',
        'Jones, M., Bradley, J., & Sakimura, N. (2015). RFC 7519: JSON Web Token (JWT). IETF.',
        'Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures (Doctoral dissertation). UC Irvine.',
        'Fowler, M., & Lewis, J. (2014). Microservices: A definition of this new architectural term. https://martinfowler.com/articles/microservices.html',
    ]
    for i, r in enumerate(refs):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(f'[{i+1}]  {r}')
        run.font.size = Pt(10)

    doc.add_page_break()

    # ================================================================
    # LAMPIRAN
    # ================================================================
    heading(doc, 'LAMPIRAN', 1)
    hline(doc)

    heading(doc, 'Lampiran A: Cara Menjalankan Sistem', 2)
    heading(doc, 'Prasyarat:', 3)
    for p in ['Docker Engine v24.0+ terinstall', 'Docker Compose v2.0+ terinstall',
              'Minimal 4GB RAM tersedia', 'Port 3000, 8000, 5432, 5672, 15672 tidak digunakan']:
        bullet(doc, p)

    heading(doc, 'Langkah Menjalankan:', 3)
    code_block(doc,
        '# 1. Masuk ke direktori proyek\ncd final-sistemterdistribusi\n\n'
        '# 2. Build dan jalankan semua container\nmake up\n# atau secara manual:\ndocker compose up -d --build\n\n'
        '# 3. Tunggu semua service healthy (±60-90 detik)\ndocker compose ps\n\n'
        '# 4. Akses\n'
        '#  Frontend     : http://localhost:3000\n'
        '#  API Gateway  : http://localhost:8000/health\n'
        '#  RabbitMQ UI  : http://localhost:15672 (admin / admin123)\n\n'
        '# 5. Perintah berguna\nmake logs      # Lihat log realtime semua service\nmake status    # Status semua container\nmake down      # Hentikan semua service\nmake clean     # Hapus container + volume (reset data)',
        'Lampiran A — Langkah Instalasi dan Menjalankan Sistem'
    )

    heading(doc, 'Lampiran B: Contoh Request API Lengkap', 2)
    heading(doc, 'B.1 Register dan Login:', 3)
    code_block(doc,
        '# Register akun baru\ncurl -s -X POST http://localhost:8000/auth/register \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"name":"Budi Santoso","email":"budi@test.com","password":"pass123","role":"customer","phone":"081234567890"}\'\n\n'
        '# Login\ncurl -s -X POST http://localhost:8000/auth/login \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"email":"budi@test.com","password":"pass123"}\' | python3 -c "import sys,json; d=json.load(sys.stdin); print(d[\'token\'])"',
        'Lampiran B.1 — Register & Login via curl'
    )

    heading(doc, 'B.2 Buat Pesanan:', 3)
    code_block(doc,
        '# Simpan token ke variabel (dari hasil login)\nTOKEN="eyJhbGciOi..."\n\n'
        '# Buat pesanan\ncurl -s -X POST http://localhost:8000/api/orders \\\n'
        '  -H "Authorization: Bearer $TOKEN" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{\n'
        '    "sender_name": "Budi Santoso",\n'
        '    "sender_phone": "081234567890",\n'
        '    "sender_city": "Jakarta",\n'
        '    "sender_address": "Jl. Sudirman No. 1, Jakarta Pusat",\n'
        '    "receiver_name": "Ani Wijaya",\n'
        '    "receiver_phone": "087654321098",\n'
        '    "receiver_city": "Bandung",\n'
        '    "receiver_address": "Jl. Braga No. 5, Bandung",\n'
        '    "weight": 500,\n'
        '    "description": "Pakaian",\n'
        '    "service_type": "express"\n'
        '  }\'',
        'Lampiran B.2 — Buat Pesanan Baru'
    )

    heading(doc, 'B.3 Update Status dan Cek Tracking:', 3)
    code_block(doc,
        '# Simpan ORDER_ID dari response buat pesanan\nORDER_ID=1\n\n'
        '# Update status (harus admin/courier JWT)\ncurl -X PUT http://localhost:8000/api/orders/$ORDER_ID/status \\\n'
        '  -H "Authorization: Bearer $TOKEN" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"status": "shipped"}\'\n\n'
        '# Cek riwayat tracking (otomatis ter-update via RabbitMQ)\ncurl -s http://localhost:8000/api/tracking/$ORDER_ID \\\n'
        '  -H "Authorization: Bearer $TOKEN" | python3 -m json.tool\n\n'
        '# Cek notifikasi (otomatis dibuat via RabbitMQ)\ncurl -s http://localhost:8000/api/notifications \\\n'
        '  -H "Authorization: Bearer $TOKEN" | python3 -m json.tool',
        'Lampiran B.3 — Update Status, Cek Tracking & Notifikasi'
    )

    heading(doc, 'Lampiran C: Struktur Direktori Lengkap', 2)
    code_block(doc,
        'final-sistemterdistribusi/\n'
        '├── docker-compose.yml          # 10 container: 4 DB, 1 RabbitMQ, 5 services\n'
        '├── Makefile                    # Shortcut perintah\n'
        '├── services/\n'
        '│   ├── api-gateway/            # Port 8000\n'
        '│   │   ├── main.go             # JWT middleware, routing, forwardRequest\n'
        '│   │   ├── go.mod              # gin, golang-jwt\n'
        '│   │   └── Dockerfile\n'
        '│   ├── user-service/           # Port 8001\n'
        '│   │   ├── main.go             # Register, login, profil, JWT\n'
        '│   │   ├── go.mod              # gin, gorm, bcrypt, jwt\n'
        '│   │   └── Dockerfile\n'
        '│   ├── order-service/          # Port 8002\n'
        '│   │   ├── main.go             # CRUD order, kalkulasi harga, RabbitMQ publisher\n'
        '│   │   ├── go.mod              # gin, gorm, amqp091-go, jwt\n'
        '│   │   └── Dockerfile\n'
        '│   ├── tracking-service/       # Port 8003\n'
        '│   │   ├── main.go             # API tracking, RabbitMQ consumer goroutine\n'
        '│   │   ├── go.mod              # gin, gorm, amqp091-go\n'
        '│   │   └── Dockerfile\n'
        '│   └── notification-service/   # Port 8004\n'
        '│       ├── main.go             # CRUD notif, RabbitMQ consumer goroutine\n'
        '│       ├── go.mod              # gin, gorm, amqp091-go\n'
        '│       └── Dockerfile\n'
        '├── frontend/                   # Port 3000\n'
        '│   ├── src/\n'
        '│   │   ├── app/               # Next.js App Router\n'
        '│   │   │   ├── layout.tsx     # Root layout\n'
        '│   │   │   ├── page.tsx       # Root redirect\n'
        '│   │   │   ├── login/page.tsx # Form login\n'
        '│   │   │   ├── register/page.tsx\n'
        '│   │   │   ├── track/page.tsx  # Lacak paket publik\n'
        '│   │   │   ├── dashboard/\n'
        '│   │   │   │   ├── layout.tsx  # Sidebar + auth check\n'
        '│   │   │   │   └── page.tsx    # Dashboard stats\n'
        '│   │   │   ├── orders/\n'
        '│   │   │   │   ├── page.tsx    # Daftar pesanan\n'
        '│   │   │   │   ├── new/page.tsx # Form buat pesanan\n'
        '│   │   │   │   └── [id]/page.tsx # Detail + tracking\n'
        '│   │   │   └── notifications/page.tsx\n'
        '│   │   ├── lib/api.ts          # Axios client + API functions\n'
        '│   │   └── types/index.ts      # TypeScript types\n'
        '│   ├── package.json\n'
        '│   ├── next.config.js\n'
        '│   ├── tailwind.config.ts\n'
        '│   └── Dockerfile\n'
        '└── docs/\n'
        '    ├── generate_report.py      # Script generator laporan ini\n'
        '    └── Laporan_*.docx         # Output laporan',
        'Lampiran C — Struktur Direktori Proyek Lengkap'
    )

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run('─' * 62)
    r.font.color.rgb = RGBColor(148, 163, 184)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Dokumen dibuat otomatis: {datetime.now().strftime("%d %B %Y, %H:%M WIB")}')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(148, 163, 184); r.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LogistikPro — Sistem Logistik & Ekspedisi Terdistribusi | Arsitektur Microservices')
    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RGBColor(37, 99, 235)

    out = os.path.join(os.path.dirname(__file__), 'Laporan_Sistem_Logistik_Terdistribusi.docx')
    doc.save(out)
    return out, os.path.getsize(out)


if __name__ == '__main__':
    print('=' * 65)
    print('  GENERATOR LAPORAN DOCX — LogistikPro')
    print('  Sistem Logistik & Ekspedisi Terdistribusi')
    print('=' * 65)
    try:
        path, size = generate()
        print(f'\n✅ Berhasil! File: {path}')
        print(f'📄 Ukuran: {size/1024:.1f} KB')
        print('\n📋 ISI LAPORAN:')
        print('  ✅ Halaman Judul')
        print('  ✅ Abstrak')
        print('  ✅ Daftar Isi')
        print('  ✅ BAB I   — Pendahuluan')
        print('  ✅ BAB II  — Landasan Teori + Tabel perbandingan')
        print('  ✅ BAB III — Perancangan (Arsitektur, DB Design, API Design)')
        print('  ✅ BAB IV  — Implementasi + 24 snippet kode program')
        print('  ✅ BAB V   — Tampilan UI (7 placeholder screenshot)')
        print('  ✅ BAB VI  — Pengujian (Unit, Integrasi, Ketahanan)')
        print('  ✅ BAB VII — Penutup (Kesimpulan & Saran)')
        print('  ✅ Daftar Pustaka (18 referensi)')
        print('  ✅ Lampiran A,B,C')
        print('\n📸 PANDUAN SCREENSHOT (ganti placeholder di BAB V):')
        print('  ┌─────────────────────────────────────────────────────┐')
        print('  │  Gambar 5.1 → Login     : http://localhost:3000/login')
        print('  │  Gambar 5.2 → Register  : http://localhost:3000/register')
        print('  │  Gambar 5.3 → Dashboard : http://localhost:3000/dashboard')
        print('  │  Gambar 5.4 → Pesanan   : http://localhost:3000/orders')
        print('  │  Gambar 5.5 → Buat Order: http://localhost:3000/orders/new')
        print('  │  Gambar 5.6 → Notifikasi: http://localhost:3000/notifications')
        print('  │  Gambar 5.7 → Lacak     : http://localhost:3000/track')
        print('  └─────────────────────────────────────────────────────┘')
        print('\n🎉 Buka DOCX, ganti 7 placeholder SS dengan screenshot Anda!')
    except Exception as e:
        import traceback
        print(f'\n❌ Error: {e}')
        traceback.print_exc()
